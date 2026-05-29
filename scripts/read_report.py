"""
报告文本提取工具
支持多语言（简体中文、繁体中文、英文、粤语白话文）、PDF/Word/Excel
自动检测语言和编码，输出 UTF-8 文本供 AI 分析
用法: python read_report.py <文件路径>
"""
import sys
import os
from pathlib import Path


def detect_language(text: str) -> dict:
    """检测文本语言和字符集（支持简/繁/英/粤/日/韩）"""
    result = {
        "language": "unknown",
        "script": "unknown",
        "encoding": "utf-8",
    }

    latin = 0
    japanese = 0
    korean = 0
    cjk_chars = 0  # 中日韩统一表意文字总数

    for ch in text:
        cp = ord(ch)
        # CJK统一表意文字（包含简繁）
        if 0x4E00 <= cp <= 0x9FFF:
            cjk_chars += 1
        elif 0x3400 <= cp <= 0x4DBF:
            cjk_chars += 1
        elif cp in range(0x3040, 0x30A0):
            japanese += 1
        elif cp in range(0xAC00, 0xD7B0):
            korean += 1
        elif cp in range(0x41, 0x5B) or cp in range(0x61, 0x7B):
            latin += 1

    # ---- 繁体中文检测 ----
    # 方法：统计"仅繁体有、简体没有"的字符出现次数
    # 这些字的Unicode编码与对应简体字不同
    traditional_only = set(
        "臺灣衛擊復盡誌雲係華麗榮風豐廠處號資訊體電畫裡這裡那裡什麼關係系統發展計畫參與對於關於"
        "為時會機對個學發來後開門頭現實力經從動長問過進實體關東陽書車馬魚龍"
        "兒見貝門頁風飛馬魚鳥鹵麥黃黑點齊齒龜"
        "說無與嗎麼誰哪唉呀吧呢哈呵噠嗎那個這為甚麼"
        "幾辦員單團場報備準則閱經聲責選擇購買賣還"
        "後門開關頭現實力經從動長問過進"
        "筆記錄層導屆幣帳號碼線網組續總結果統計約"
        "軟體運算資料網路電腦資訊瀏覽器"
        "醫藥護療診證試檢查驗"
        "鐵銀銅鋼錶鏡鐘鍵針鋒鎖鍊"
        "夠夠嗎嗎嗎嗎嗎"
    )
    # 去除重复
    traditional_only = set(traditional_only)

    trad_count = 0
    for ch in traditional_only:
        trad_count += text.count(ch)

    # 繁体中文文本中常见的独有字比率一般在 0.5%~3%
    trad_ratio = trad_count / max(cjk_chars, 1)

    # ---- 粤语白话文检测 ----
    cantonese_chars = set("嘅咗哋嚟喺啱嘢冇冧佢哋噃啩喎乜嘥瞓")
    canto_count = sum(text.count(ch) for ch in cantonese_chars)
    canto_ratio = canto_count / max(cjk_chars, 1)

    # ---- 日文检测 ----
    # 日文假名 + CJK
    kana_count = sum(1 for ch in text if 0x3040 <= ord(ch) <= 0x30FF)

    # ---- 判定逻辑 ----
    if cjk_chars == 0 and latin > 100:
        result["language"] = "en"
        result["script"] = "English"
    elif canto_ratio > 0.005:
        result["language"] = "yue"
        result["script"] = "Cantonese Vernacular (粤语白话文)"
    elif trad_ratio > 0.005 and cjk_chars > 100:
        result["language"] = "zh-Hant"
        result["script"] = "Traditional Chinese (繁体中文)"
    elif cjk_chars > latin * 2:
        result["language"] = "zh-Hans"
        result["script"] = "Simplified Chinese (简体中文)"
    elif kana_count > cjk_chars * 0.3:
        result["language"] = "ja"
        result["script"] = "Japanese (日文)"
    elif korean > cjk_chars * 0.3:
        result["language"] = "ko"
        result["script"] = "Korean (韩文)"
    elif latin > cjk_chars * 3:
        result["language"] = "en"
        result["script"] = "English"

    # 附加：检测是否混合（如中英对照报告）
    if latin > cjk_chars * 0.1 and cjk_chars > 100:
        result["mixed"] = True
        result["note"] = "文本包含中英混合内容"

    return result


def extract_pdf(filepath: str) -> str:
    """提取PDF文本，自动处理编码"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("[ERROR] 请安装 PyMuPDF: pip install PyMuPDF")
        sys.exit(1)

    doc = fitz.open(filepath)
    pages_text = []
    font_info = {}

    for i in range(doc.page_count):
        page = doc[i]

        # 收集字体信息
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        fname = span.get("font", "")
                        if fname and fname not in font_info:
                            font_info[fname] = span.get("size", 0)

        # 提取文本（PyMuPDF内部已是UTF-8）
        text = page.get_text()
        pages_text.append(text)

    doc.close()
    full_text = "\n".join(pages_text)

    # 检测语言
    sample = full_text[:5000]
    lang = detect_language(sample)

    return {
        "text": full_text,
        "pages": len(pages_text),
        "language": lang,
        "fonts": list(font_info.keys()),
    }


def extract_docx(filepath: str) -> dict:
    """提取Word文档文本"""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] 请安装 python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # 也提取表格
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                tables_text.append(" | ".join(row_text))

    full_text = "\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n--- 表格 ---\n" + "\n".join(tables_text)

    lang = detect_language(full_text[:5000])

    return {
        "text": full_text,
        "pages": len(doc.paragraphs),
        "language": lang,
        "fonts": [],
    }


def extract_xlsx(filepath: str) -> dict:
    """提取Excel表格文本"""
    try:
        import openpyxl
    except ImportError:
        print("[ERROR] 请安装 openpyxl: pip install openpyxl")
        sys.exit(1)

    wb = openpyxl.load_workbook(filepath, data_only=True)
    all_text = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_text.append(f"\n=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = [str(c) if c is not None else "" for c in row]
            if any(row_text):
                all_text.append(" | ".join(row_text))

    full_text = "\n".join(all_text)
    lang = detect_language(full_text[:5000])

    return {
        "text": full_text,
        "pages": len(wb.sheetnames),
        "language": lang,
        "fonts": [],
    }


def extract_txt(filepath: str) -> dict:
    """读取纯文本文件，自动检测编码"""
    # 尝试多种编码
    encodings = ["utf-8", "utf-16", "gbk", "gb2312", "big5", "latin-1", "cp1252"]
    text = None

    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                text = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if text is None:
        # 最后尝试 chardet
        try:
            import chardet
            with open(filepath, "rb") as f:
                raw = f.read()
            detected = chardet.detect(raw)
            text = raw.decode(detected["encoding"])
        except ImportError:
            text = f"[ERROR] 无法检测文件编码，请手动转为UTF-8: {filepath}"

    lang = detect_language(text[:5000] if text else "")

    return {
        "text": text or "",
        "pages": 1,
        "language": lang,
        "fonts": [],
    }


def main():
    if len(sys.argv) < 2:
        print("用法: python read_report.py <报告文件路径> [输出路径]")
        print("支持格式: PDF, DOCX, XLSX, TXT")
        print("")
        print("输出:")
        print("  1. 控制台打印报告摘要（语言、页数、字体）")
        print("  2. 文本内容保存为 UTF-8 文件")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"[ERROR] 文件不存在: {filepath}")
        sys.exit(1)

    ext = Path(filepath).suffix.lower()

    print(f"[INFO] 正在读取: {filepath}")
    print(f"[INFO] 文件格式: {ext}")

    # 提取文本
    if ext == ".pdf":
        result = extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        result = extract_docx(filepath)
    elif ext in (".xlsx", ".xls"):
        result = extract_xlsx(filepath)
    elif ext in (".txt", ".md", ".csv"):
        result = extract_txt(filepath)
    else:
        print(f"[ERROR] 不支持的格式: {ext}")
        sys.exit(1)

    # 打印摘要
    lang = result["language"]
    print(f"[INFO] 总页数/Sheet数: {result['pages']}")
    print(f"[INFO] 总字符数: {len(result['text'])}")
    print(f"[INFO] 检测语言: {lang['language']} ({lang['script']})")
    if result.get("fonts"):
        print(f"[INFO] 检测到的字体: {result['fonts'][:10]}")

    # 确定输出路径
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        base = os.path.splitext(filepath)[0]
        output_path = f"{base}_extracted.txt"

    # 写入UTF-8文件
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(result["text"])

    print(f"[OK] 文本已保存至: {output_path}")
    print(f"[INFO] 编码: UTF-8 | 可直接用 Read 工具或文本编辑器打开")


if __name__ == "__main__":
    main()